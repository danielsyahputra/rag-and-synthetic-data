import argparse
import json
import re
import time

from docx import Document
from docx2pdf import convert
from docx.shared import Inches, Mm
from docxtpl import DocxTemplate, InlineImage
from wrapper import OpenAIWrapper


def parse_arguments():
    parser = argparse.ArgumentParser(description="Book Generator")
    parser.add_argument("--topic", type=str, required=True, help="Book topic")
    parser.add_argument(
        "--target-audience", type=str, required=True, help="Target audience"
    )
    parser.add_argument(
        "--num-chapters", type=int, default=1, help="Number of chapters"
    )
    parser.add_argument(
        "--num-subsections",
        type=int,
        default=5,
        help="Number of subsections per chapter",
    )
    parser.add_argument(
        "--output-docx", type=str, default="tmp/temp.docx", help="Output DOCX file path"
    )
    parser.add_argument(
        "--output-pdf", type=str, default="tmp/book.pdf", help="Output PDF file path"
    )
    parser.add_argument("--book-template", type=str, help="Book template file path")
    parser.add_argument("--cover-template", type=str, help="Cover template file path")
    parser.add_argument(
        "--preview", action="store_true", help="Generate preview version"
    )
    parser.add_argument("--preview-image", type=str, help="Preview image path")
    return parser.parse_args()


class BookGenerator:
    def __init__(self) -> None:
        self.wrapper = OpenAIWrapper()

    def generate_title(self, topic, target_audience):
        convo_id = self.wrapper.start_convo(
            system="You are a book author with 20+ experience."
        )
        title_prompt = (
            f'Buat judul menarik untuk buku tentang "{topic}". '
            f'Target pembaca: "{target_audience}". '
            'Maksimal 9 kata, berikan "janji besar" yang memikat pembaca. '
            "Berikan dalam Bahasa Indonesia."
        )

        title = self.wrapper.msg_in_convo(convo_id, title_prompt)
        title = title.replace('"', "")
        return title

    def verify_outline(self, outline, num_chapters, num_subsections):
        if len(outline.items()) != num_chapters:
            return False
        for chapter, subtopics in outline.items():
            if len(chapter) < 15:
                return False
            if type(subtopics) != list:
                return False
            if len(subtopics) != num_subsections:
                return False
        return True

    def generate_outline(
        self, topic, target_audience, title, num_chapters, num_subsections
    ):
        convo_id = self.wrapper.start_convo(
            "You are a book author with 20+ experience."
        )
        outline_prompt = (
            f'Kami sedang menulis buku elektronik berjudul "{title}". Buku ini tentang '
            f'"{topic}". Pembaca target kami adalah: "{target_audience}". Buatlah '
            f"kerangka menyeluruh untuk buku ini yang akan memiliki {num_chapters} bab. "
            f"Setiap bab harus memiliki tepat {num_subsections} subbab. Format output "
            "untuk prompt: dictionary Python dengan key: judul bab, value: sebuah list/array "
            "yang berisi judul-judul subbab dalam bab tersebut (subtopik harus berada "
            'dalam list). Judul bab harus diawali dengan nomor bab, seperti ini: "Bab 5: '
            '[judul bab]". Judul subbab harus diawali dengan {nomor bab.nomor subbab}, '
            'seperti ini: "5.4: [judul subbab]". Berikan jawaban dalam Bahasa Indonesia.'
        )

        outline_json = self.wrapper.msg_in_convo(convo_id, outline_prompt)
        outline_json = outline_json[
            outline_json.find("{") : outline_json.rfind("}") + 1
        ]
        outline = json.loads(outline_json)
        if not self.verify_outline(outline, num_chapters, num_subsections):
            raise Exception("Outline not well formed!")
        return outline

    def generate_chapter_content(
        self, title, topic, target_audience, idx, chapter, subtopic
    ):
        convo_id = self.wrapper.start_convo(
            "You are a book author with 20+ experience."
        )
        num_words_str = "500 to 700"
        content_prompt = (
            f'Kami sedang menulis buku elektronik berjudul "{title}". '
            f'Secara keseluruhan, buku ini tentang "{topic}". '
            f'Pembaca target kami adalah: "{target_audience}". '
            f'Kami sedang menulis bagian #{idx+1} untuk bab: "{chapter}". '
            f'Gunakan minimal {num_words_str} kata untuk menulis konten lengkap mengenai subtopik: "{subtopic}". '
            "Konten harus semaksimal mungkin membantu pembaca. "
            "Sertakan fakta kuantitatif dan statistik, lengkap dengan referensi. "
            "Jelaskan secara mendalam sesuai kebutuhan. "
            "Anda bisa membaginya menjadi beberapa paragraf jika diperlukan. "
            "Konten harus ditulis dalam bentuk paragraf yang padu. "
            'Jangan menyertakan bagian "[Masukkan ___]" yang akan memerlukan pengeditan manual dalam buku nantinya. '
            'Jika Anda merasa perlu menulis "masukkan [kosong]" di manapun, jangan lakukan itu (ini sangat penting). '
            "Jika Anda tidak mengetahui sesuatu, jangan masukkan dalam konten. "
            "Jangan sertakan informasi tambahan seperti jumlah kata, karena seluruh konten akan langsung masuk ke dalam buku elektronik untuk pembaca, tanpa proses editing manusia. "
            f"Ingat minimum {num_words_str} kata, harap patuhi itu. "
            "Berikan jawaban dalam Bahasa Indonesia."
        )

        content = self.wrapper.msg_in_convo(convo_id, content_prompt)

        def remove_subtopic_from_content(content, subtopic, search_limit=200):
            pattern = r"\d\.\d"

            match = re.search(pattern, content[:search_limit])

            if match:
                newline_index = content.find("\n", match.end())

                if newline_index != -1:
                    result_string = content[newline_index + 1 :]
                    content = content.strip()
                    return result_string
            return content

        try:
            content = content.strip()
            content = remove_subtopic_from_content(content, subtopic)
        except:
            pass
        return content

    def generate_docs(
        self,
        topic,
        target_audience,
        title,
        outline,
        docx_file,
        book_template,
        preview,
        actionable_steps=False,
    ):
        document = Document(book_template)
        document.add_page_break()
        document.add_heading("Table of Contents")

        for chapter, subtopics in outline.items():
            document.add_heading(chapter, level=2)
            for idx, subtopic in enumerate(subtopics):
                document.add_heading("\t" + subtopic, level=3)

        document.add_page_break()
        chapter_num = 1

        for chapter, subtopics in outline.items():
            document.add_heading(chapter, level=1)

            intro_content = self.generate_chapter_content(
                title, topic, target_audience, 0, chapter, "Introduction"
            )
            time.sleep(1)
            document.add_paragraph(intro_content)

            try:
                viz_path = self.generate_visualization(
                    title, intro_content, chapter_num
                )
                document.add_picture(viz_path, width=Inches(6))
                document.add_paragraph(
                    "Figure " + str(chapter_num) + ": Chapter Overview Visualization"
                )
            except Exception as e:
                print(
                    f"Warning: Could not generate visualization for chapter {chapter_num}: {e}"
                )

            subtopics_content = []
            for idx, subtopic in enumerate(subtopics):
                if preview and idx >= 2:
                    break

                document.add_heading(subtopic, level=2)
                content = self.generate_chapter_content(
                    title, topic, target_audience, idx, chapter, subtopic
                )
                time.sleep(1)
                document.add_paragraph(content)

                if len(content) > 300:
                    try:
                        viz_path = self.generate_visualization(
                            title, content, f"{chapter_num}_{idx+1}"
                        )
                        document.add_picture(viz_path, width=Inches(5))
                        document.add_paragraph(
                            f"Figure {chapter_num}.{idx+1}: {subtopic} Visualization"
                        )
                    except Exception as e:
                        print(
                            f"Warning: Could not generate visualization for subtopic {subtopic}: {e}"
                        )

                subtopics_content.append(content)

            if preview:
                document.add_heading(
                    "Preview Completed - Purchase Full Book To Read More!"
                )
                break

            if chapter_num < len(outline.items()):
                document.add_page_break()
            chapter_num += 1

            try:
                import os

                for file in os.listdir("tmp"):
                    if file.startswith(f"chapter_{chapter_num}") and file.endswith(
                        ".png"
                    ):
                        os.remove(os.path.join("tmp", file))
            except Exception as e:
                print(f"Warning: Could not clean up visualization files: {e}")

        document.save(docx_file)

    def generate_cover(
        self, cover_template, title, topic, target_audience, output_file, preview
    ):
        temp_docx = "tmp/temp_cover_intermediate.docx"
        temp_photo = "tmp/preview_photo.png"

        output_pdf = output_file.rsplit(".", 1)[0] + ".pdf"

        doc = DocxTemplate(cover_template)

        if preview:
            image_path = preview
        else:
            self.generate_cover_photo(title, topic, target_audience, temp_photo)
            time.sleep(1)
            image_path = temp_photo

        try:
            imagen = InlineImage(doc, image_path, width=Mm(120))
            context = {"title": title, "subtext": "DNL Publishing", "image": imagen}
            doc.render(context)

            doc.save(temp_docx)

            convert(temp_docx, output_pdf)

        finally:
            import os

            try:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                if os.path.exists(temp_photo) and not preview:
                    os.remove(temp_photo)
            except Exception as e:
                print(f"Warning: Could not clean up temporary files: {e}")

    def generate_cover_photo(self, title, topic, target_audience, img_output):
        convo_id = self.wrapper.start_convo(
            "You are a book author with 20+ experience."
        )
        cover_prompt = (
            f'Kami memiliki buku elektronik berjudul {title}. Buku ini tentang "{topic}". '
            f'Pembaca kami adalah: "{target_audience}". Tuliskan deskripsi singkat dan '
            "lugas tentang foto yang akan digunakan pada sampul buku. Jangan menyebutkan "
            "sampul atau foto dalam jawaban Anda. Contohnya, jika judulnya "
            '"Cara Diet untuk Wanita Paruh Baya", jawaban yang tepat adalah '
            '"wanita paruh baya sedang berolahraga"'
        )

        print(cover_prompt)
        dalle_prompt = self.wrapper.msg_in_convo(convo_id, cover_prompt)
        print(dalle_prompt)

        img_data = self.wrapper.generate_photo(dalle_prompt)
        with open(img_output, "wb") as handler:
            handler.write(img_data)

    def generate_visualization(self, title, content, chapter_num):
        """
        Generates a relevant visualization based on the chapter content.
        Returns the image in PNG format.
        """
        convo_id = self.wrapper.start_convo(
            "You are a data visualization expert with 20+ years of experience."
        )
        viz_prompt = (
            f'Dari konten bab "{title}": "{content[:500]}..." '
            "buat prompt DALL-E untuk ilustrasi profesional. "
            "Fokus pada diagram bisnis, bagan alur, atau grafik konseptual, "
            "bukan foto realistis."
        )

        dalle_prompt = self.wrapper.msg_in_convo(convo_id, viz_prompt)
        img_data = self.wrapper.generate_photo(dalle_prompt)

        # Save the visualization
        img_path = f"tmp/chapter_{chapter_num}_viz.png"
        with open(img_path, "wb") as handler:
            handler.write(img_data)

        return img_path


if __name__ == "__main__":
    args = parse_arguments()
    generator = BookGenerator()

    title = generator.generate_title(
        topic=args.topic, target_audience=args.target_audience
    )
    print("Title: ", title)

    outline = generator.generate_outline(
        topic=args.topic,
        target_audience=args.target_audience,
        title=title,
        num_chapters=args.num_chapters,
        num_subsections=args.num_subsections,
    )
    print("Outline: ", outline)

    generator.generate_docs(
        topic=args.topic,
        target_audience=args.target_audience,
        title=title,
        outline=outline,
        docx_file=args.output_docx,
        book_template=args.book_template,
        preview=args.preview,
    )

    cover_image_path = "tmp/image.png"
    generator.generate_cover_photo(
        title=title,
        topic=args.topic,
        target_audience=args.target_audience,
        img_output=cover_image_path,
    )

    generator.generate_cover(
        cover_template=args.cover_template or args.output_docx,
        title=title,
        topic=args.topic,
        target_audience=args.target_audience,
        output_file=args.output_pdf,
        preview=args.preview_image or cover_image_path,
    )
